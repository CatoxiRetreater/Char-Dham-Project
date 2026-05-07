"""Build End-Term Report - Part 2: Chapters 1-3"""
import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = Path(__file__).parent
ASSETS = BASE / "report_assets"
OUT = BASE / "End_Term_Report.docx"
with open(ASSETS/'metrics.json') as f: M = json.load(f)
doc = Document(str(OUT))

style = doc.styles['Normal']
style.font.name = 'Times New Roman'; style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

def heading(txt, level=1):
    h = doc.add_heading(txt, level=level)
    for r in h.runs: r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)
def para(txt, bold=False, italic=False, size=12):
    p = doc.add_paragraph()
    r = p.add_run(txt); r.font.name='Times New Roman'; r.font.size=Pt(size)
    r.bold=bold; r.italic=italic; p.paragraph_format.line_spacing=1.5
    return p
def img(name, w=5.5, caption=None):
    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    path = ASSETS/name
    if path.exists(): p.add_run().add_picture(str(path), width=Inches(w))
    if caption:
        c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER
        cr=c.add_run(caption); cr.font.name='Times New Roman'; cr.font.size=Pt(10); cr.italic=True
def pagebreak(): doc.add_page_break()
def add_table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers), style='Light Grid Accent 1')
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h
        for p in c.paragraphs:
            for r in p.runs: r.bold=True; r.font.size=Pt(10); r.font.name='Times New Roman'
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(10); r.font.name='Times New Roman'

# ==================== CHAPTER 1: INTRODUCTION ====================
heading("CHAPTER 1", level=1)
heading("INTRODUCTION", level=1)

heading("1.1 Background and Context", level=2)
para("The Char Dham pilgrimage circuit, nestled within the Garhwal Himalayas of Uttarakhand in northern India, constitutes one of the most sacred and heavily visited religious tourism corridors in the world. Comprising four ancient shrines, namely Kedarnath at an altitude of 3,583 metres in the Rudraprayag district, Badrinath at 3,133 metres in the Chamoli district, Gangotri at 3,100 metres in the Uttarkashi district, and Yamunotri at 3,293 metres also in Uttarkashi, this circuit holds immense spiritual significance for millions of Hindu devotees who undertake the arduous journey each year between the months of May and November when the mountain passes remain accessible.")
para("In recent decades, the confluence of improved road infrastructure, government-sponsored pilgrimage facilitation programmes such as the Char Dham Highway Development Project, rising disposable incomes across the Indian middle class, and the proliferation of social media promoting religious tourism has resulted in a dramatic escalation of annual visitor volumes. Government records from the Uttarakhand Tourism Development Board indicate that cumulative annual pilgrim footfall across the four shrines has risen from approximately 1.5 million in 2010 to over 5.6 million by 2023, representing a compounded annual growth rate of roughly ten to twelve per cent. During peak months of May, June, September, and October, daily visitor counts at individual shrines regularly exceed their designated ecological carrying capacities by margins of thirty to fifty per cent.")
para("This unprecedented surge in human activity within ecologically fragile high-altitude zones has precipitated a cascade of environmental consequences that threaten the long-term sustainability of both the pilgrimage tradition and the Himalayan ecosystem that sustains it. Documented impacts include accelerated deforestation for road widening and accommodation construction along approach routes, measurable declines in Normalized Difference Vegetation Index values indicating vegetation stress in shrine peripheries, solid waste accumulation that overwhelms municipal collection infrastructure, contamination of glacial meltwater streams that serve as tributaries to the Ganges and Yamuna river systems, increased landslide susceptibility due to slope destabilisation from construction activity, and disruption of alpine and sub-alpine biodiversity corridors.")

img('chart_annual_trends.png', 5.5, 'Figure 1.1: Annual Pilgrim Footfall Trends Across Char Dham Shrines')

heading("1.2 Problem Statement", level=2)
para("The existing approach to tourism management in the Char Dham region is predominantly reactive in nature. Authorities typically respond to overcrowding, waste accumulation, or infrastructure failures after these problems have already manifested, rather than anticipating and preventing them through data-driven forecasting and proactive intervention. There is no integrated decision-support platform that combines real-time meteorological monitoring, historical tourism pattern analysis, satellite-based vegetation health assessment, and predictive machine learning models into a coherent analytical framework.")
para("This absence of an integrated, AI-powered monitoring system creates several critical operational gaps. First, tourism authorities lack the ability to predict pilgrim volumes with sufficient lead time to implement crowd management strategies such as staggered entry permits or alternate routing. Second, environmental monitoring agencies have no real-time quantitative metric that captures the composite environmental pressure across multiple stress dimensions simultaneously. Third, there is no mechanism to simulate hypothetical scenarios, such as the impact of a twenty per cent increase in tourist volume or a significant temperature anomaly, on ecosystem health before these conditions materialise. Fourth, the disconnect between tourism data, climate data, and vegetation health data prevents holistic analysis of the tourism-environment nexus.")
para("The core problem this project addresses can therefore be stated as follows: there exists no integrated, data-driven, AI-powered decision-support system that can monitor, analyse, predict, and visualise the complex interplay between tourism pressure, climatic variability, and ecosystem health across the Char Dham pilgrimage region in a manner that enables proactive rather than reactive management of environmental carrying capacity violations.")

heading("1.3 Motivation", level=2)
para("Being from India, we have personally witnessed the transformation that pilgrimage tourism brings to hill regions, especially during peak yatra seasons. Places like Kedarnath, Badrinath, Gangotri, and Yamunotri become extremely crowded within a span of a few weeks, and the environmental consequences are visible to the naked eye. Mountains of discarded plastic bottles and food packaging line the trekking routes, temporary accommodation structures mushroom without adequate waste disposal infrastructure, and the pristine mountain streams that pilgrims revere take on a murky quality from upstream contamination.")
para("During and after the yatra season, issues such as overflowing garbage dumps, water pollution from untreated sewage, tree cutting for road expansion, and increased landslide risks along destabilised slopes become clearly noticeable. These problems are not just environmental; they represent a direct threat to pilgrim safety, as the 2013 Kedarnath disaster tragically demonstrated when flash floods and landslides killed thousands of pilgrims and devastated the shrine infrastructure.")
para("This motivated us to explore whether modern artificial intelligence and data science techniques could be applied to predict environmental pressure before it reaches critical levels. With the availability of structured tourism data from government sources, climate records from meteorological agencies, and satellite imagery from programmes like MODIS and Landsat, we recognised that the necessary data infrastructure already existed. What was missing was an analytical platform capable of synthesising these disparate data streams into actionable intelligence for decision-makers.")

heading("1.4 Objectives", level=2)
para("The primary objective of this project is to develop an AI-based ecosystem intelligence system that can analyse and predict tourism pressure and ecosystem stress in the Char Dham region using historical and real-time data. The specific objectives are enumerated below:")
objs = [
    "To collect, clean, integrate, and preprocess historical tourism footfall and climate datasets spanning the period 2010 to 2024 for all four Char Dham shrines, creating a unified analytical dataset with derived temporal and environmental features.",
    "To design and implement a composite Ecosystem Stress Index that quantifies environmental pressure on a normalised zero-to-hundred scale by aggregating weighted sub-indices for pilgrim capacity utilization, temperature anomaly, and rainfall deviation.",
    "To develop and evaluate three complementary machine learning models, namely Random Forest Regression, Seasonal ARIMA, and Long Short-Term Memory neural networks, for forecasting monthly pilgrim counts with quantified prediction uncertainty.",
    "To integrate real-time weather data from the OpenWeatherMap API and satellite-derived NDVI vegetation health indices from MODIS products into the monitoring pipeline with graceful fallback mechanisms.",
    "To implement a multi-factor alert engine that generates severity-graded risk assessments across four categories: tourist overload, climate stress, NDVI degradation, and composite ESI thresholds.",
    "To build a what-if scenario simulator that allows stakeholders to explore the impact of hypothetical changes in pilgrim volume, temperature, rainfall, and vegetation health on ecosystem stress through interactive sensitivity analysis.",
    "To provide model transparency through SHAP-based explainability that reveals feature contribution patterns driving individual predictions.",
    "To deploy the complete system as an interactive, enterprise-grade Streamlit dashboard with six dedicated analytical pages, interactive Folium-based geospatial maps, and downloadable PDF and CSV reports.",
]
for i,o in enumerate(objs, 1):
    p = doc.add_paragraph()
    r = p.add_run(f"{i}. {o}"); r.font.name='Times New Roman'; r.font.size=Pt(12)
    p.paragraph_format.line_spacing = 1.5

heading("1.5 Scope of the Project", level=2)
para("The scope of this project encompasses the complete software lifecycle from data ingestion through analytical modelling to interactive visualisation and alert generation. The system operates on two primary structured datasets, the Tourist Footfall Dataset and the Climate Dataset, covering four shrines across a fifteen-year temporal span from 2010 to 2024. The spatial scope is limited to the four Char Dham shrine zones and their immediate surrounding regions within approximately a five-kilometre radius for satellite data analysis.")
para("The system delivers the following major functional capabilities: automated data integration and preprocessing with temporal feature engineering; composite Ecosystem Stress Index computation; interactive exploratory data analysis with correlation matrices, scatter plots, and seasonality heatmaps; three-model predictive forecasting with confidence intervals; SHAP-based model explainability; real-time weather monitoring with API integration; geospatial intelligence with NDVI heatmaps, LULC classification, and interactive shrine mapping; multi-factor alert generation with configurable thresholds; what-if scenario simulation with sensitivity analysis; and downloadable reports in PDF and CSV formats.")
para("Explicit exclusions from the current scope include: mobile application development, real-time IoT sensor integration, user authentication and role-based access control, relational database backend, Hindi language localisation, and direct integration with government administrative systems. These items are identified as future enhancement opportunities in Chapter 9 of this report.")

heading("1.6 Target Beneficiaries", level=2)
para("The system is designed to serve four primary user groups. Tourism Authorities such as the Uttarakhand Tourism Development Board require predictive footfall data and stress-level alerts to regulate pilgrim access, coordinate logistics, and enforce carrying capacity limits during peak seasons. Environmental Monitoring Agencies such as the Uttarakhand Environment Protection and Pollution Control Board need quantitative ESI metrics and NDVI trend visualisations to assess anthropogenic impact on vegetation and water resources. Research Scholars and Academic Institutions studying climate-tourism interactions and Himalayan ecology benefit from the analytical modules, correlation analyses, and model comparison frameworks. Policy Makers and Government Administrators require high-level decision support summaries, scenario-based forecasts, and AI-generated recommendations to formulate evidence-based sustainable tourism policies.")

heading("1.7 Project Timeline", level=2)
para("The project was executed across five distinct stages over a fourteen-month development period from January 2025 to May 2026. Stage 1 (January to March 2025) encompassed data collection, cleaning, and initial exploratory analysis. Stage 2 (April to June 2025) focused on satellite data integration, NDVI computation, and LULC classification. Stage 3 (July to October 2025) involved model development including ARIMA, LSTM, Random Forest, and the ESI formulation. Stage 4 (November 2025 to February 2026) covered dashboard development, UI design, API integration, and the alert engine. Stage 5 (March to May 2026) comprised real-time data integration, testing, optimisation, and final deployment.")

heading("1.8 Report Organization", level=2)
para("This report is organised into ten chapters. Chapter 1 provides the introduction, problem statement, and objectives. Chapter 2 presents a comprehensive literature review establishing the theoretical foundation. Chapter 3 details the system analysis and design including UML diagrams and architecture. Chapter 4 describes the datasets and preprocessing pipeline. Chapter 5 explains the methodology and algorithms employed. Chapter 6 discusses the implementation details. Chapter 7 presents the results and output screens. Chapter 8 covers testing and validation. Chapter 9 discusses limitations and future enhancements. Chapter 10 concludes the report with a summary of achievements and significance.")
pagebreak()

# ==================== CHAPTER 2: LITERATURE REVIEW ====================
heading("CHAPTER 2", level=1)
heading("LITERATURE REVIEW", level=1)

heading("2.1 Tourism Impact on Himalayan Ecosystems", level=2)
para("The environmental impact of mass tourism on fragile mountain ecosystems has been a subject of growing academic scrutiny over the past two decades. Kuniyal et al. (2004) conducted one of the earliest systematic assessments of solid waste generation along pilgrimage routes in the Indian Himalayas, documenting that daily waste output during peak season exceeded the capacity of existing municipal collection systems by a factor of three to five. Their findings established a direct linear correlation between daily visitor counts and solid waste volumes, a relationship that our system operationalises through the Estimated Waste Tons feature derived as a function of pilgrim count.")
para("Sati and Kumar (2015) examined the ecological carrying capacity of pilgrimage sites in the Garhwal Himalayas, proposing threshold-based management frameworks that closely parallel the capacity utilization ratio employed in our Ecosystem Stress Index formulation. Their research highlighted that sites exceeding seventy-five per cent of ecological carrying capacity for sustained periods exhibited measurable vegetation decline within a two-to-three year lag period. This temporal relationship between overcrowding and ecological degradation provides empirical justification for the proactive alert thresholds implemented in our multi-factor alert engine.")
para("Negi and Maikhuri (2017) analysed the socio-economic and environmental dimensions of the Char Dham Yatra specifically, identifying seasonal pilgrim concentration as the primary driver of forest cover loss, soil compaction, and water contamination in shrine peripheries. Their work provided the domain-specific context that informed our selection of shrine-level granularity for data collection and analysis, as well as the inclusion of district-level spatial metadata in both datasets.")

heading("2.2 Remote Sensing and NDVI Analysis", level=2)
para("The Normalized Difference Vegetation Index, first proposed by Rouse et al. (1974) and operationalised through satellite-borne multispectral sensors, has become the standard metric for quantifying vegetation health and density from space. NDVI is computed as the normalised difference between near-infrared and visible red reflectance bands, yielding values between negative one and positive one, where values above 0.5 indicate dense healthy vegetation and values below 0.2 suggest barren or stressed landscapes.")
para("Piao et al. (2011) demonstrated the utility of MODIS-derived NDVI time series for detecting long-term vegetation trends across the Tibetan Plateau, a biogeographic region sharing significant climatic and ecological characteristics with the Char Dham zone. Their methodology of computing monthly NDVI averages over five-kilometre buffer zones around points of interest directly inspired our geospatial module's approach to NDVI extraction using both Google Earth Engine and the ORNL MODIS REST API.")
para("Roy et al. (2019) applied Landsat-based NDVI analysis to assess forest cover changes along pilgrimage corridors in Uttarakhand, finding statistically significant negative correlations between proximity to major shrine access roads and NDVI values. Their spatial analysis framework influenced our implementation of distance-weighted NDVI heatmaps in the geospatial page, where vegetation density is visualised as a function of distance from shrine centres.")

heading("2.3 Time Series Forecasting in Tourism", level=2)
para("Tourism demand forecasting has evolved from simple extrapolation methods to sophisticated statistical and deep learning approaches. Song and Li (2008) provided a comprehensive review of tourism forecasting methodologies, establishing that ARIMA-family models consistently outperformed naive benchmarks for short-to-medium term seasonal tourism prediction, while neural network approaches showed advantages for capturing nonlinear patterns in longer horizons.")
para("Box and Jenkins (1976) formulated the foundational ARIMA methodology that our system extends to the seasonal variant SARIMA. The seasonal component is critical for pilgrimage tourism data, which exhibits strong twelve-month periodicity driven by the fixed annual yatra schedule and monsoon weather patterns. Our implementation employs a two-phase grid search optimisation across sixty-four coarse parameter combinations followed by nine refined fits to identify the optimal order and seasonal order that minimise the Akaike Information Criterion.")
para("Hochreiter and Schmidhuber (1997) introduced the Long Short-Term Memory architecture that addresses the vanishing gradient problem inherent in traditional recurrent neural networks. Subsequent applications by Cankurt and Subasi (2016) demonstrated LSTM efficacy for tourism arrivals prediction, achieving superior performance to traditional statistical methods on datasets with complex temporal dependencies. Our LSTM implementation uses a two-layer architecture with sixty-four and thirty-two units respectively, incorporating dropout regularisation and early stopping to prevent overfitting on the relatively small shrine-level time series.")

heading("2.4 Composite Stress Indices in Environmental Science", level=2)
para("The concept of composite environmental indices that aggregate multiple stress dimensions into a single interpretable metric has a rich precedent in ecological science. The Environmental Stress Index proposed by Moran et al. (2001) for heat stress assessment in occupational health demonstrated the utility of weighted linear combinations of normalised physiological variables for classification into actionable risk categories. Our Ecosystem Stress Index adopts a structurally similar approach, combining normalised capacity utilization, temperature anomaly, and rainfall deviation through empirically calibrated weights.")
para("Palmer (1965) introduced the Palmer Drought Severity Index, which became one of the most widely adopted composite environmental metrics, demonstrating that multi-variable indices with clearly defined threshold-based classification schemes could effectively communicate complex environmental states to non-technical stakeholders. This principle of threshold-based classification with intuitive colour coding, mapping continuous index values to discrete categories such as Low, Moderate, and Critical, directly informed the design of our ESI classification system and its integration into the alert engine.")

heading("2.5 Dashboard and Decision Support Systems", level=2)
para("The emergence of open-source data science frameworks has democratised the development of interactive analytical dashboards. Streamlit, introduced by Snowflake Inc. in 2019, provides a Python-native framework for building data applications with minimal frontend development overhead. Its reactive programming model, where the entire script re-executes upon user interaction, simplifies state management but necessitates careful use of caching decorators to maintain acceptable performance with computationally expensive operations such as model training.")
para("Plotly and its Python interface provide publication-quality interactive visualisations with hover tooltips, zoom capabilities, and responsive layouts that are essential for exploratory data analysis in a dashboard context. Folium, which wraps the Leaflet.js mapping library, enables integration of interactive geospatial visualisations including choropleth maps, heatmaps, and marker-based overlays within Python-generated web applications.")

heading("2.6 Research Gap and Contribution", level=2)
para("Despite the substantial body of literature on tourism impact assessment, NDVI-based vegetation monitoring, and time series forecasting individually, there remains a significant gap in integrated systems that combine all three analytical dimensions within a unified, interactive decision-support platform specifically designed for pilgrimage tourism management in ecologically sensitive mountain regions.")
para("The key contributions of this project that address this research gap are: first, the development of a novel composite Ecosystem Stress Index that quantifies multi-dimensional environmental pressure at shrine-level granularity; second, the implementation of a three-model predictive framework with standardised evaluation metrics that enables comparative assessment of statistical, ensemble, and deep learning approaches for pilgrim count forecasting; third, the integration of real-time weather data and satellite-derived vegetation indices with historical tourism analytics in a single platform; fourth, the provision of what-if scenario simulation capability with sensitivity analysis that enables stakeholders to explore policy implications before implementation; and fifth, the deployment of all analytical capabilities within an accessible, enterprise-grade web dashboard with automated report generation.")

add_table(
    ['Author(s)', 'Year', 'Focus Area', 'Key Finding', 'Gap Addressed'],
    [
        ['Kuniyal et al.', '2004', 'Waste in Himalayan tourism', 'Waste exceeds capacity 3-5x in peak season', 'No predictive framework'],
        ['Sati & Kumar', '2015', 'Carrying capacity', '75% threshold leads to vegetation decline', 'No real-time monitoring'],
        ['Negi & Maikhuri', '2017', 'Char Dham impact', 'Seasonal concentration is primary driver', 'No integrated AI system'],
        ['Piao et al.', '2011', 'MODIS NDVI trends', 'Effective for long-term vegetation monitoring', 'Not applied to tourism'],
        ['Roy et al.', '2019', 'Forest cover change', 'Negative correlation near shrine roads', 'No dashboard integration'],
        ['Song & Li', '2008', 'Tourism forecasting', 'ARIMA best for seasonal short-term', 'Single model approach'],
        ['Box & Jenkins', '1976', 'ARIMA methodology', 'Foundation for time series analysis', 'No seasonal tourism use'],
        ['Hochreiter et al.', '1997', 'LSTM networks', 'Solves vanishing gradient problem', 'Not applied to pilgrimage'],
    ]
)
para("Table 2.1: Summary of Key Literature and Research Gaps", italic=True, size=10)
pagebreak()

# ==================== CHAPTER 3: SYSTEM ANALYSIS & DESIGN ====================
heading("CHAPTER 3", level=1)
heading("SYSTEM ANALYSIS AND DESIGN", level=1)

heading("3.1 Requirement Analysis", level=2)
para("The system requirements were elicited through a combination of domain literature review, analysis of existing government tourism management practices, consultation with the project guide, and iterative refinement during the development lifecycle. Requirements are classified into functional and non-functional categories as specified in the Software Requirements Specification document prepared during the initial project phase.")

heading("3.1.1 Functional Requirements", level=3)
add_table(
    ['ID', 'Requirement', 'Priority', 'Module'],
    [
        ['FR-01', 'Load and merge Tourist Footfall and Climate Excel datasets', 'High', 'Data Loader'],
        ['FR-02', 'Compute Ecosystem Stress Index for each shrine-month record', 'High', 'Feature Engine'],
        ['FR-03', 'Display real-time weather from OpenWeatherMap API', 'High', 'Weather API'],
        ['FR-04', 'Train Random Forest model with 19 climate+tourism features', 'High', 'Models'],
        ['FR-05', 'Train SARIMA with optimized seasonal order selection', 'High', 'Models'],
        ['FR-06', 'Train LSTM neural network for sequence prediction', 'Medium', 'Models'],
        ['FR-07', 'Generate multi-factor risk alerts with severity levels', 'High', 'Alerts'],
        ['FR-08', 'Display interactive Folium maps with shrine status', 'Medium', 'Geospatial'],
        ['FR-09', 'Compute and display NDVI vegetation health metrics', 'Medium', 'Geospatial'],
        ['FR-10', 'Provide what-if scenario simulator with sensitivity analysis', 'Medium', 'Simulator'],
        ['FR-11', 'Generate downloadable PDF and CSV reports', 'Low', 'Report Gen'],
        ['FR-12', 'Display SHAP explainability for RF predictions', 'Medium', 'Models'],
    ]
)
para("Table 3.1: Functional Requirements Specification", italic=True, size=10)

heading("3.1.2 Non-Functional Requirements", level=3)
add_table(
    ['ID', 'Requirement', 'Metric'],
    [
        ['NFR-01', 'Dashboard initial load time', '< 5 seconds on standard hardware'],
        ['NFR-02', 'Model training time (RF)', '< 3 seconds for 200 trees'],
        ['NFR-03', 'API response caching TTL', '30 minutes to respect rate limits'],
        ['NFR-04', 'Browser compatibility', 'Chrome 90+, Firefox 85+, Edge 90+'],
        ['NFR-05', 'Minimum display resolution', '1280 x 720 pixels'],
        ['NFR-06', 'Memory footprint', '< 500 MB RAM under normal operation'],
        ['NFR-07', 'Reproducibility', 'Fixed random seeds (42) across all stochastic operations'],
        ['NFR-08', 'Graceful degradation', 'Fallback to historical data when APIs unavailable'],
    ]
)
para("Table 3.2: Non-Functional Requirements", italic=True, size=10)

heading("3.2 SWOT Analysis", level=2)
add_table(
    ['Category', 'Description'],
    [
        ['Strengths', 'Multi-model ML approach with comparative evaluation; Real-time API integration with caching; Composite ESI metric; Interactive visualizations; SHAP explainability; What-if simulator'],
        ['Weaknesses', 'File-based storage without RDBMS; No user authentication; Limited to four shrines; NDVI relies on synthetic fallback when GEE unavailable'],
        ['Opportunities', 'IoT sensor integration for real-time crowd monitoring; Mobile app deployment; Government partnership for official deployment; Extension to other pilgrimage circuits'],
        ['Threats', 'API rate limits and service disruptions; Data quality issues in government records; Climate model uncertainty; Evolving pilgrimage patterns post-COVID'],
    ]
)
para("Table 3.3: SWOT Analysis of the Char Dham Intelligence Dashboard", italic=True, size=10)

heading("3.3 System Architecture", level=2)
para("The system follows a modular layered architecture comprising four principal tiers: the Data Layer responsible for ingestion, caching, and preprocessing; the Analytics Layer housing the feature engineering, ML models, and alert engine; the Integration Layer managing external API connections; and the Presentation Layer delivering the interactive Streamlit dashboard. This separation of concerns ensures maintainability, testability, and extensibility of individual components without impacting the overall system stability.")
img('chart_architecture.png', 6.0, 'Figure 3.1: System Architecture Diagram')
para("The architecture diagram above illustrates the data flow from raw Excel datasets and external APIs through the processing pipeline to the six interactive dashboard pages. The core package contains seven Python modules, each encapsulating a distinct functional domain: data_loader.py for data pipeline management, weather_api.py for OpenWeatherMap integration, feature_engine.py for ESI and TPI computation, models.py for ML training and prediction, alerts.py for multi-factor risk assessment, geospatial.py for map generation and NDVI analysis, and report_generator.py for PDF and CSV export.")

heading("3.4 UML Diagrams", level=2)
heading("3.4.1 Use Case Diagram", level=3)
para("The system supports four primary actor categories: Tourism Authority users who access monitoring, prediction, and alert functionalities; Environmental Analysts who utilise geospatial and NDVI analysis tools; Researchers who leverage the EDA, model comparison, and SHAP explainability modules; and the System itself which acts as an autonomous actor for scheduled data refresh, model retraining, and alert generation. Key use cases include View Dashboard Overview, Monitor Real-Time Weather, Generate Pilgrim Forecast, Configure Alert Thresholds, Run What-If Simulation, Export PDF Report, and Analyse NDVI Trends.")

heading("3.4.2 Architecture Diagrams", level=3)
img('chart_lstm_arch.png', 5.5, 'Figure 3.4: LSTM Neural Network Architecture')
para("The LSTM architecture implements a sequential model with an input layer accepting twelve-month lookback windows, followed by an LSTM layer with sixty-four units and return sequences enabled, a dropout layer with twenty per cent probability, a second LSTM layer with thirty-two units, another dropout layer, a dense layer with sixteen units and ReLU activation, and a final dense output layer with a single unit for point prediction.")

heading("3.5 Data Flow Diagram", level=2)
para("The Level-0 data flow diagram represents the system as a single process receiving inputs from three external entities: the Excel Dataset Files providing historical tourism and climate data, the OpenWeatherMap API providing real-time meteorological data, and the MODIS Satellite Service providing NDVI vegetation indices. The system produces outputs to two external entities: the Dashboard User Interface rendering interactive visualisations, and the Report Files generating downloadable PDF and CSV exports.")
para("At Level-1, the central process decomposes into six sub-processes: Data Ingestion and Merging, Feature Engineering and ESI Computation, ML Model Training and Prediction, Alert Generation, Geospatial Processing, and Visualisation Rendering. Data flows between these sub-processes through pandas DataFrame objects maintained in Streamlit session state, providing efficient in-memory data sharing across dashboard pages without redundant computation.")

heading("3.6 Technology Stack", level=2)
add_table(
    ['Component', 'Technology', 'Version', 'Purpose'],
    [
        ['Language', 'Python', '3.10+', 'Core development language'],
        ['Web Framework', 'Streamlit', '>=1.30.0', 'Interactive dashboard UI'],
        ['ML Framework', 'scikit-learn', '>=1.3.0', 'Random Forest, preprocessing'],
        ['Deep Learning', 'TensorFlow/Keras', '>=2.15.0', 'LSTM neural network'],
        ['Time Series', 'statsmodels', '>=0.14.0', 'SARIMA model'],
        ['Explainability', 'SHAP', '>=0.44.0', 'Feature contribution analysis'],
        ['Visualization', 'Plotly', '>=5.18.0', 'Interactive charts'],
        ['Mapping', 'Folium', '>=0.15.0', 'Geospatial maps'],
        ['Data Processing', 'Pandas + NumPy', '>=2.0 / >=1.24', 'Data manipulation'],
        ['HTTP Client', 'Requests', '>=2.31.0', 'API communication'],
        ['PDF Generation', 'fpdf2', '>=2.7.0', 'Report export'],
        ['Environment', 'python-dotenv', '>=1.0.0', 'API key management'],
    ]
)
para("Table 3.4: Technology Stack and Dependencies", italic=True, size=10)
pagebreak()

doc.save(str(OUT))
print(f"Part 2 done: Chapters 1-3 added. Saved to {OUT}")
