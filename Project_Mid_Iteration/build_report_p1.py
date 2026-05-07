"""Build End-Term Report DOCX - Part 1: Setup + Front Matter + Ch1-2"""
import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).parent
ASSETS = BASE / "report_assets"
OUT = BASE / "End_Term_Report.docx"
with open(ASSETS/'metrics.json') as f: M = json.load(f)

doc = Document()

# --- Page setup ---
for sec in doc.sections:
    sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17); sec.right_margin = Cm(2.54)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'; style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

def heading(txt, level=1):
    h = doc.add_heading(txt, level=level)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0,0,0)
    return h

def para(txt, bold=False, italic=False, align=None, size=12):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size)
    r.bold = bold; r.italic = italic
    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right': p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = 1.5
    return p

def img(name, w=5.5, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    path = ASSETS / name
    if path.exists(): r.add_picture(str(path), width=Inches(w))
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = c.add_run(caption)
        cr.font.name = 'Times New Roman'; cr.font.size = Pt(10); cr.italic = True

def pagebreak():
    doc.add_page_break()

# ====================== TITLE PAGE ======================
for _ in range(4): doc.add_paragraph()
para("AI-Based Tourism Pressure and Ecosystem Stress\nManagement Dashboard for Char Dham Region", bold=True, align='center', size=18)
doc.add_paragraph()
para("A\nProject Report\nsubmitted in partial fulfillment of the\nrequirements for the award of the degree of", align='center', size=12)
doc.add_paragraph()
para("BACHELOR OF TECHNOLOGY\nin\nCOMPUTER SCIENCE & ENGINEERING", bold=True, align='center', size=14)
doc.add_paragraph()
para("by", align='center')
doc.add_paragraph()
para("Sangam Khanna (SAP: 500123465 | Roll No: R2142230152) — Batch 4, Data Science\nManan Nasa (SAP: 500123471 | Roll No: R2142230154) — Batch 12, AIML", align='center', size=11)
doc.add_paragraph()
para("under the guidance of\nDr. Sachi Choudhary", bold=True, align='center')
doc.add_paragraph()
para("School of Computer Science\nUniversity of Petroleum & Energy Studies\nBidholi, Via Prem Nagar, Dehradun, Uttarakhand\nMay 2026", align='center', size=11)
pagebreak()

# ====================== DECLARATION ======================
heading("CANDIDATE'S DECLARATION", level=1)
para("We hereby certify that the project work entitled \"AI-Based Tourism Pressure and Ecosystem Stress Management Dashboard for Char Dham Region\" in partial fulfillment of the requirements for the award of the Degree of BACHELOR OF TECHNOLOGY in Computer Science & Engineering, submitted to the School of Computer Science, University of Petroleum & Energy Studies, Dehradun, is an authentic record of our own work carried out during the period of January 2025 to May 2026 under the supervision of Dr. Sachi Choudhary.")
para("The matter presented in this project report has not been submitted by us for the award of any other degree of this or any other University.")
doc.add_paragraph()
para("Sangam Khanna\nSAP ID: 500123465 | Roll No: R2142230152")
para("Manan Nasa\nSAP ID: 500123471 | Roll No: R2142230154")
doc.add_paragraph()
para("This is to certify that the above statement made by the candidates is correct to the best of my knowledge.")
doc.add_paragraph()
para("Date: _____________ May 2026\t\t\t\t\tDr. Sachi Choudhary\n\t\t\t\t\t\t\t\t\t\tProject Guide")
pagebreak()

# ====================== ACKNOWLEDGEMENT ======================
heading("ACKNOWLEDGEMENT", level=1)
para("We wish to express our deep gratitude to our guide Dr. Sachi Choudhary, for all the advice, encouragement, and constant support she has given us throughout our project work. Her expert guidance and constructive criticism at every stage of this project were instrumental in shaping the final outcome. Without her mentorship, the successful completion of this project would not have been possible.")
para("We sincerely thank our respected Prof. Rajiv Nandwani, Head of the Department of Computer Science, for his great support in facilitating our project within the School of Computer Science at UPES. His leadership and vision for academic excellence have been a source of inspiration.")
para("We are also grateful to the Dean, School of Computer Science, UPES, for giving us the necessary facilities, infrastructure, and computational resources to carry out our project work successfully. We also thank our Course Coordinator for the continuous encouragement and academic support provided throughout the semester.")
para("We would like to thank all our friends, classmates, and peers for their help, feedback, and constructive criticism during our project work. Their insights during brainstorming sessions and their willingness to test our dashboard provided valuable perspectives that improved the quality of our system.")
para("Finally, we have no words to express our sincere gratitude towards our parents and family members, whose unwavering support, patience, and encouragement throughout our academic journey have been the foundation upon which all our achievements rest.")
doc.add_paragraph()
para("Sangam Khanna\nManan Nasa", align='right')
para("May 2026, Dehradun", align='right')
pagebreak()

# ====================== ABSTRACT ======================
heading("ABSTRACT", level=1)
para("The Char Dham pilgrimage circuit in Uttarakhand, comprising Kedarnath, Badrinath, Gangotri, and Yamunotri, attracts millions of pilgrims annually, placing immense pressure on fragile Himalayan ecosystems. The unregulated surge in tourist footfall, particularly during peak yatra seasons spanning May through October, has led to measurable environmental degradation including vegetation loss, soil erosion, waste accumulation, and infrastructure strain that exceeds the ecological carrying capacities of these high-altitude shrine zones.")
para("This project presents the design, development, and deployment of an AI-powered Ecosystem Intelligence Dashboard that integrates historical tourism data, real-time meteorological feeds, satellite-derived vegetation indices, and machine learning forecasting models into a unified decision-support platform. The system operates on two curated datasets, a Tourist Footfall Dataset containing monthly pilgrim counts across four shrines from 2010 to 2024, and a Climate Dataset with meteorological variables including temperature, rainfall, humidity, wind speed, and solar radiation for the same temporal and spatial scope.")
para("The analytical core of the system implements three complementary predictive models. A Random Forest Regressor trained on nineteen climate and tourism features achieves an R-squared score of {:.3f} with five-fold cross-validation mean R-squared of {:.3f}. A SARIMA model with optimized seasonal order captures monthly periodicity for twelve-month-ahead forecasting. A Long Short-Term Memory neural network with a two-layer architecture processes sequential patterns in the pilgrim count time series. All three models are compared using standardized metrics including RMSE, MAE, R-squared, and MAPE.".format(M['rf_r2'], M['rf_cv_mean']))
para("The system introduces a novel composite Ecosystem Stress Index that quantifies environmental pressure on a normalized zero-to-hundred scale by combining weighted sub-indices for capacity utilization, temperature anomaly, and rainfall deviation. A multi-factor alert engine generates color-coded risk assessments across four severity levels. The geospatial module integrates NDVI vegetation health data from MODIS satellite products through a cascading fallback architecture spanning Google Earth Engine, ORNL MODIS REST API, and historical baseline models.")
para("The dashboard is deployed as an interactive Streamlit web application with six distinct pages encompassing overview analytics, real-time monitoring with OpenWeatherMap API integration, multi-model predictions with SHAP explainability, geospatial intelligence with Folium-based interactive maps, a configurable alert center with PDF and CSV export capabilities, and a what-if scenario simulator with sensitivity analysis. The system represents a significant advancement over existing reactive tourism management approaches by providing proactive, data-driven decision support for sustainable pilgrimage tourism in ecologically sensitive Himalayan regions.")
para("Keywords: Ecosystem Stress Index, Char Dham, Random Forest, SARIMA, LSTM, NDVI, Tourism Pressure, Streamlit Dashboard, SHAP Explainability, Carrying Capacity", italic=True)
pagebreak()

# ====================== TABLE OF CONTENTS ======================
heading("TABLE OF CONTENTS", level=1)
toc_items = [
    ("","Candidate's Declaration","ii"),("","Acknowledgement","iii"),("","Abstract","iv"),
    ("","Table of Contents","vi"),("","List of Figures","viii"),("","List of Tables","x"),
    ("1","Introduction","1"),("1.1","Background and Context","1"),("1.2","Problem Statement","3"),
    ("1.3","Motivation","4"),("1.4","Objectives","5"),("1.5","Scope of the Project","6"),
    ("1.6","Target Beneficiaries","7"),("1.7","Project Timeline","8"),("1.8","Report Organization","9"),
    ("2","Literature Review","10"),("2.1","Tourism Impact on Himalayan Ecosystems","10"),
    ("2.2","Remote Sensing and NDVI Analysis","12"),("2.3","Time Series Forecasting in Tourism","13"),
    ("2.4","Composite Stress Indices","14"),("2.5","Dashboard and Decision Support Systems","15"),
    ("2.6","Research Gap and Contribution","16"),
    ("3","System Analysis and Design","18"),("3.1","Requirement Analysis","18"),
    ("3.2","SWOT Analysis","20"),("3.3","System Architecture","21"),
    ("3.4","UML Diagrams","23"),("3.5","Data Flow Diagrams","26"),("3.6","Technology Stack","27"),
    ("4","Dataset Description and Preprocessing","29"),("4.1","Data Sources","29"),
    ("4.2","Tourist Footfall Dataset","30"),("4.3","Climate Dataset","32"),
    ("4.4","Data Integration Pipeline","33"),("4.5","Data Cleaning","34"),
    ("4.6","Feature Engineering","35"),("4.7","Exploratory Data Analysis","37"),
    ("5","Methodology and Algorithms","41"),("5.1","Ecosystem Stress Index","41"),
    ("5.2","Tourism Pressure Index","43"),("5.3","Random Forest Regressor","44"),
    ("5.4","SARIMA Model","46"),("5.5","LSTM Neural Network","48"),
    ("5.6","SHAP Explainability","50"),("5.7","Multi-Factor Alert Engine","51"),
    ("5.8","NDVI Analysis Pipeline","52"),
    ("6","Implementation","54"),("6.1","Development Environment","54"),
    ("6.2","Project Structure","55"),("6.3","Core Modules","56"),
    ("6.4","API Integration","58"),("6.5","UI/UX Design","59"),("6.6","Deployment","60"),
    ("7","Results and Output Screens","61"),("7.1","Dashboard Home Page","61"),
    ("7.2","Real-Time Monitoring","63"),("7.3","Predictions Module","64"),
    ("7.4","Geospatial Intelligence","66"),("7.5","Alert Center","67"),
    ("7.6","What-If Simulator","68"),("7.7","Model Performance Summary","69"),
    ("8","Testing and Validation","71"),("8.1","Testing Strategy","71"),
    ("8.2","Model Validation Results","72"),("8.3","Edge Case Handling","73"),
    ("9","Limitations and Future Enhancements","74"),("9.1","Current Limitations","74"),
    ("9.2","Future Enhancements","75"),
    ("10","Conclusion","77"),
    ("","References","79"),("","Appendix A: Glossary","81"),("","Appendix B: Dataset Schemas","82"),
]
for num, title, pg in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    prefix = f"{num}\t" if num else ""
    r = p.add_run(f"{prefix}{title}")
    r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    if not num or (num and '.' not in num): r.bold = True
    tab = p.add_run(f"\t{pg}")
    tab.font.name = 'Times New Roman'; tab.font.size = Pt(11)
pagebreak()

# ====================== LIST OF FIGURES ======================
heading("LIST OF FIGURES", level=1)
figs = [
    ("1.1","Annual Pilgrim Footfall Trends Across Char Dham Shrines","2"),
    ("1.2","Project Timeline and Development Stages","8"),
    ("3.1","System Architecture Diagram","22"),
    ("3.2","Use Case Diagram","24"),("3.3","Activity Diagram","25"),
    ("3.4","LSTM Neural Network Architecture","26"),
    ("4.1","Feature Correlation Matrix","37"),("4.2","Seasonality Heatmap: Kedarnath","38"),
    ("4.3","Climate vs Pilgrim Count Scatter Analysis","39"),
    ("4.4","Distribution of Key Features","39"),("4.5","Box Plots by Shrine","40"),
    ("5.1","ESI Computation Pipeline","42"),("5.2","Sensitivity Analysis","53"),
    ("7.1","Current Capacity Utilization by Shrine","62"),
    ("7.2","ESI Trends Across All Shrines","63"),("7.3","ESI Comparison Bar Chart","64"),
    ("7.4","Random Forest: Actual vs Predicted and Feature Importance","65"),
    ("7.5","RF Prediction Residual Distribution","65"),
    ("7.6","SARIMA 12-Month Forecast with Confidence Intervals","66"),
    ("7.7","STL Seasonal Decomposition","66"),
    ("7.8","Model Performance Comparison","69"),
    ("7.9","Year-over-Year Pilgrim Growth Rate","70"),
    ("7.10","Monthly Average Pilgrim Distribution","70"),
    ("7.11","Tourism-Waste Correlation Analysis","70"),
]
for num, title, pg in figs:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(f"Fig. {num}\t{title}\t\t{pg}")
    r.font.name = 'Times New Roman'; r.font.size = Pt(11)
pagebreak()

# ====================== LIST OF TABLES ======================
heading("LIST OF TABLES", level=1)
tables = [
    ("2.1","Literature Review Summary","16"),("3.1","Functional Requirements","19"),
    ("3.2","Non-Functional Requirements","20"),("3.3","SWOT Analysis","21"),
    ("3.4","Technology Stack","28"),("4.1","Tourist Footfall Dataset Schema","31"),
    ("4.2","Climate Dataset Schema","32"),("4.3","Descriptive Statistics by Shrine","36"),
    ("5.1","ESI Weight Configuration","42"),("5.2","Alert Level Classification","51"),
    ("5.3","Random Forest Hyperparameters","45"),
    ("7.1","Cross-Model Performance Comparison","69"),
    ("8.1","Model Validation Summary","72"),
]
for num, title, pg in tables:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(f"Table {num}\t{title}\t\t{pg}")
    r.font.name = 'Times New Roman'; r.font.size = Pt(11)
pagebreak()

# Save intermediate
doc.save(str(OUT))
print(f"Part 1 done: Front matter saved to {OUT}")
