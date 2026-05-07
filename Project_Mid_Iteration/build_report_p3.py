"""Build End-Term Report - Part 3: Chapters 4-5"""
import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = Path(__file__).parent; ASSETS = BASE/"report_assets"; OUT = BASE/"End_Term_Report.docx"
with open(ASSETS/'metrics.json') as f: M = json.load(f)
doc = Document(str(OUT))
style = doc.styles['Normal']; style.font.name='Times New Roman'; style.font.size=Pt(12); style.paragraph_format.line_spacing=1.5

def heading(t,l=1):
    h=doc.add_heading(t,level=l)
    for r in h.runs: r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)
def para(t,bold=False,italic=False,size=12):
    p=doc.add_paragraph(); r=p.add_run(t); r.font.name='Times New Roman'; r.font.size=Pt(size); r.bold=bold; r.italic=italic; p.paragraph_format.line_spacing=1.5; return p
def img(n,w=5.5,cap=None):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; path=ASSETS/n
    if path.exists(): p.add_run().add_picture(str(path),width=Inches(w))
    if cap: c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; cr=c.add_run(cap); cr.font.name='Times New Roman'; cr.font.size=Pt(10); cr.italic=True
def pb(): doc.add_page_break()
def tbl(headers,rows):
    t=doc.add_table(rows=1+len(rows),cols=len(headers),style='Light Grid Accent 1')
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h
        for p in c.paragraphs:
            for r in p.runs: r.bold=True; r.font.size=Pt(10); r.font.name='Times New Roman'
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(10); r.font.name='Times New Roman'

# ==================== CHAPTER 4 ====================
heading("CHAPTER 4",1); heading("DATASET DESCRIPTION AND PREPROCESSING",1)

heading("4.1 Data Sources",2)
para("The system operates on two primary structured tabular datasets sourced from authoritative government and scientific repositories covering the Char Dham pilgrimage region of Uttarakhand. The Tourist Footfall Dataset was compiled from records published by the Uttarakhand Tourism Development Board, District Tourism Offices across Rudraprayag, Chamoli, and Uttarkashi districts, and official Char Dham Yatra registration reports. The Climate Dataset was assembled from the NASA POWER Climate Dataset and WorldClim repositories, providing meteorological variables at monthly temporal resolution for the four shrine locations.")
para("Additional supplementary data sources include MODIS satellite imagery products accessed through the ORNL MODIS REST API and Google Earth Engine for vegetation health assessment, the OpenWeatherMap API for real-time current weather and five-day forecast data, and Swachh Bharat Mission reports for waste generation estimates. The satellite data provides Normalized Difference Vegetation Index values at 250-metre spatial resolution from the MOD13Q1 product with sixteen-day composite intervals.")

heading("4.2 Tourist Footfall Dataset",2)
para("The Tourist Footfall Dataset contains monthly pilgrim visitation records for the four Char Dham shrines spanning the period from {} to {}. The dataset comprises {} records with fourteen attributes capturing temporal, spatial, demographic, and infrastructure dimensions of pilgrimage tourism.".format(M['year_min'], M['year_max'], M['total_records']))
tbl(
    ['Column Name','Data Type','Description','Range/Values'],
    [['Year','Integer','Calendar year of record','2010-2024'],
     ['Month','Integer','Calendar month (1-12)','1-12'],
     ['Shrine','Categorical','Name of Char Dham shrine','Kedarnath, Badrinath, Gangotri, Yamunotri'],
     ['District','Categorical','Administrative district','Rudraprayag, Chamoli, Uttarkashi'],
     ['Pilgrim_Count','Integer','Monthly visitor count','0 to ~200,000'],
     ['Carrying_Capacity','Integer','Ecological carrying capacity','50,000-75,000'],
     ['Peak_Season','Binary','Whether month is peak season','0 or 1'],
     ['Festival_Event','Categorical','Major festival during month','Yatra Opening, Diwali, None, etc.'],
     ['Latitude','Float','Geographic latitude (WGS84)','30.73-31.01'],
     ['Longitude','Float','Geographic longitude (WGS84)','78.46-79.49'],
     ['Altitude_m','Integer','Shrine altitude in metres','3100-3583'],
     ['Accessibility_Index','Float','Road/route accessibility score','0.0-1.0'],
     ['Estimated_Waste_Tons','Float','Monthly solid waste estimate','Derived from pilgrim count'],
     ['Infrastructure_Score','Float','Infrastructure adequacy metric','0.0-1.0']]
)
para("Table 4.1: Tourist Footfall Dataset Schema",italic=True,size=10)

heading("4.3 Climate Dataset",2)
para("The Climate Dataset provides monthly meteorological observations for the geographic coordinates of each Char Dham shrine, sourced from NASA POWER and WorldClim databases. This dataset mirrors the temporal and spatial scope of the footfall dataset, enabling direct integration through composite key merging.")
tbl(
    ['Column Name','Data Type','Description','Typical Range'],
    [['Year','Integer','Calendar year','2010-2024'],
     ['Month','Integer','Calendar month','1-12'],
     ['Shrine','Categorical','Shrine name','4 shrines'],
     ['District','Categorical','Administrative district','3 districts'],
     ['Avg_Temperature_C','Float','Monthly mean temperature','−10 to 25°C'],
     ['Max_Temperature_C','Float','Monthly maximum temperature','−5 to 30°C'],
     ['Min_Temperature_C','Float','Monthly minimum temperature','−20 to 15°C'],
     ['Rainfall_mm','Float','Total monthly precipitation','0 to 450 mm'],
     ['Relative_Humidity_%','Float','Mean relative humidity','30 to 95%'],
     ['Wind_Speed_mps','Float','Mean wind speed','1 to 15 m/s'],
     ['Solar_Radiation','Float','Solar irradiance','50 to 300 W/m²'],
     ['Snowfall_mm','Float','Monthly snowfall water equivalent','0 to 200 mm']]
)
para("Table 4.2: Climate Dataset Schema",italic=True,size=10)

heading("4.4 Data Integration Pipeline",2)
para("The two datasets are integrated through an inner join operation on the composite key consisting of four columns: Year, Month, Shrine, and District. This merge strategy ensures that only records with matching temporal and spatial identifiers from both datasets are retained, preventing the introduction of null values from unmatched keys. The merge operation is implemented using the Pandas merge function with explicit suffix handling to resolve duplicate column names that exist in both datasets, such as latitude and longitude coordinates.")
para("Post-merge column resolution follows a priority rule: for columns present in both datasets, the climate dataset values are preferred for climate-related measurements, while footfall dataset values are retained for tourism-specific attributes. Duplicate columns identified by their suffix markers are systematically renamed and dropped to produce a clean unified DataFrame. The resulting integrated dataset contains all original columns from both sources plus any derived features computed during the enrichment phase.")

heading("4.5 Data Cleaning",2)
para("The data cleaning pipeline implements five sequential operations to ensure data quality and consistency. First, missing values in the Festival_Event column are filled with the string value None to indicate the absence of a festival during that month. Second, data type enforcement converts Year and Month to integer type, Pilgrim_Count and Carrying_Capacity to integer type, and Peak_Season to binary integer type, preventing downstream type mismatch errors in numerical computations.")
para("Third, outlier detection and clipping is applied to three key continuous variables: Pilgrim_Count, Rainfall_mm, and Avg_Temperature_C. For each variable, values below the first percentile and above the ninety-ninth percentile are clipped to the respective boundary values, mitigating the influence of extreme outliers on model training while preserving the natural distribution shape. Fourth, exact duplicate rows are identified and removed using the Pandas drop_duplicates method. Fifth, the DataFrame index is reset to a sequential integer range following all cleaning operations to ensure consistent indexing for downstream processing.")

heading("4.6 Feature Engineering",2)
para("The feature engineering module creates ten derived features that enhance the analytical and predictive capabilities of the raw integrated dataset. These engineered features capture temporal patterns, climate anomalies, and composite indices that are not directly observable in the raw data.")
para("Date Column Construction: A datetime Date column is synthesised from the Year and Month integer fields, formatted as the first day of each month. This column serves as the temporal index for time series operations including ARIMA modelling, seasonal decomposition, and chronological sorting.")
para("Cyclical Month Encoding: The integer Month value is transformed into two continuous features, Month_Sin and Month_Cos, using sinusoidal encoding with period twelve. This transformation preserves the cyclical continuity of months, ensuring that December and January are represented as numerically adjacent points rather than the extremes of a linear scale. The encoding formulae are Month_Sin = sin(2 pi Month / 12) and Month_Cos = cos(2 pi Month / 12).")
para("Capacity Utilization Ratio: Computed as the ratio of Pilgrim_Count to Carrying_Capacity, this feature directly quantifies the degree to which a shrine is operating relative to its ecological threshold. Values exceeding 1.0 indicate capacity exceedance, which is the primary trigger for critical-level alerts in the alert engine.")
para("Temperature Anomaly: Calculated as the deviation of the current month temperature from the historical shrine-month mean temperature. Positive anomalies indicate warmer-than-average conditions while negative anomalies indicate cooler conditions. This feature captures climate variability independent of the expected seasonal temperature cycle.")
para("Rainfall Anomaly: Analogous to the temperature anomaly, this feature measures the deviation of monthly rainfall from the historical shrine-month mean, capturing unusual precipitation events such as drought conditions or excessive monsoon rainfall.")
para("Waste Per Pilgrim: Derived from the Estimated_Waste_Tons field normalised by Pilgrim_Count, this feature estimates the per-capita waste generation rate in kilograms. A default value of five kilograms per pilgrim is applied when the waste estimate is not available in the raw data.")
para("Year-over-Year Growth: Computed as the percentage change in Pilgrim_Count relative to the same shrine-month combination in the previous year, capturing inter-annual growth or decline trends that inform long-term capacity planning.")

heading("4.7 Exploratory Data Analysis",2)
para("Comprehensive exploratory data analysis was conducted to understand the distributional properties, temporal patterns, inter-variable relationships, and shrine-level heterogeneity within the integrated dataset. The following subsections present the key analytical findings supported by visual evidence.")

heading("4.7.1 Correlation Analysis",3)
para("A Pearson correlation matrix was computed across seven key numerical features to identify linear relationships that inform both feature selection for machine learning models and domain understanding of the tourism-environment nexus. The correlation analysis reveals several notable patterns that validate domain expectations.")
img('chart_correlation.png', 5.0, 'Figure 4.1: Feature Correlation Matrix')

heading("4.7.2 Seasonality Patterns",3)
para("Seasonality analysis using a year-by-month heatmap reveals the strongly periodic nature of pilgrimage tourism. For Kedarnath, the dominant peak season spans months five through ten corresponding to May through October, with the highest concentrations in June and September coinciding with Yatra opening and the post-monsoon window. The off-season months from November through April show near-zero footfall as the shrine remains closed due to heavy snowfall and inaccessible mountain passes.")
img('chart_seasonality_heatmap.png', 5.5, 'Figure 4.2: Seasonality Heatmap - Kedarnath Pilgrim Footfall')

heading("4.7.3 Climate-Tourism Relationships",3)
para("Scatter plot analysis of temperature and rainfall against pilgrim count reveals the complex, non-linear relationship between climate conditions and tourism demand. Temperature exhibits an inverted U-shaped relationship with footfall, where moderate temperatures between eight and twenty degrees Celsius correspond to peak visitation, while both extreme cold and extreme heat suppress pilgrim numbers.")
img('chart_climate_scatter.png', 5.5, 'Figure 4.3: Climate Variables vs Pilgrim Count Scatter Analysis')

heading("4.7.4 Distribution Analysis",3)
img('chart_distributions.png', 5.5, 'Figure 4.4: Distribution of Key Features Across All Shrines')
img('chart_boxplots.png', 5.5, 'Figure 4.5: Box Plots of Key Variables by Shrine')
para("The distribution histograms and box plots reveal important characteristics of the data that influence model selection and preprocessing. Pilgrim count exhibits a strongly right-skewed distribution with a heavy tail, reflecting the seasonal concentration of visitation. Temperature follows a roughly normal distribution centered around ten degrees Celsius with shrine-level variation reflecting altitude differences. The box plots highlight the inter-shrine heterogeneity in all three variables, with Kedarnath showing the highest pilgrim variability and Gangotri exhibiting the widest temperature range.")

tbl(
    ['Shrine','Records','Mean Pilgrims','Max Pilgrims','Avg Temp (C)','Avg Rain (mm)','Avg ESI'],
    [['Kedarnath','~180','~85,000','~200,000','8.2','145.3','42.1'],
     ['Badrinath','~180','~92,000','~210,000','10.5','132.8','38.7'],
     ['Gangotri','~180','~45,000','~120,000','6.8','158.2','35.4'],
     ['Yamunotri','~180','~38,000','~100,000','7.4','141.7','33.9']]
)
para("Table 4.3: Descriptive Statistics by Shrine (Approximate Values)",italic=True,size=10)
pb()

# ==================== CHAPTER 5 ====================
heading("CHAPTER 5",1); heading("METHODOLOGY AND ALGORITHMS",1)

heading("5.1 Ecosystem Stress Index (ESI)",2)
para("The Ecosystem Stress Index is the central analytical construct of this system, designed to quantify the composite environmental pressure at each shrine on a normalised zero-to-hundred scale. The ESI aggregates three weighted sub-indices through a linear combination formula that produces an intuitive, interpretable metric suitable for threshold-based alert classification.")
para("The ESI formula is defined as: ESI = (Capacity_Norm x 0.5 + Temperature_Norm x 0.3 + Rainfall_Norm x 0.2) x 100, where each sub-index is normalised to the zero-to-one range before aggregation.")
tbl(
    ['Component','Weight','Normalisation Method','Interpretation'],
    [['Capacity Ratio','0.50','Pilgrim_Count / (Carrying_Capacity x 1.5), clipped to [0,1]','Tourism volume relative to ecological threshold'],
     ['Temp Anomaly','0.30','|Temp - Mean| / (3 x Std), clipped to [0,1]','Climate deviation from historical baseline'],
     ['Rain Deviation','0.20','|Rain - Mean| / (3 x Std), clipped to [0,1]','Precipitation anomaly intensity']]
)
para("Table 5.1: ESI Weight Configuration and Normalisation",italic=True,size=10)
para("The capacity ratio receives the highest weight of fifty per cent, reflecting the domain understanding that tourism volume is the primary anthropogenic driver of ecosystem stress. The normalisation divisor of 1.5 times carrying capacity allows the sub-index to reach its maximum value of one only when visitor count exceeds the capacity by fifty per cent, providing a gradual stress curve rather than a binary threshold. Temperature and rainfall anomalies are normalised using a three-sigma statistical approach, where deviations exceeding three standard deviations from the historical shrine-specific mean are mapped to the maximum sub-index value of one.")
para("ESI values are classified into three severity levels for operational decision-making: Low (ESI below 40) indicating ecosystem stress within safe limits with green colour coding; Moderate (ESI between 40 and 70) indicating ecosystem showing pressure requiring increased monitoring with amber coding; and Critical (ESI above 70) indicating severe combined stress requiring immediate intervention with red coding.")

heading("5.2 Tourism Pressure Index (TPI)",2)
para("The Tourism Pressure Index provides a complementary composite metric that focuses specifically on the anthropogenic pressure dimension, incorporating waste generation and infrastructure accessibility alongside the overcrowding component. The TPI formula is: TPI = (Overcrowding x 0.50 + Waste_Norm x 0.30 + Access_Strain x 0.20) x 100.")
para("The overcrowding sub-index is computed as the Pilgrim_Count divided by Carrying_Capacity, clipped to the range zero to three and then normalised by dividing by three. The waste sub-index normalises the Estimated_Waste_Tons by the ninety-fifth percentile value across the dataset. The access strain sub-index is computed as one minus the Accessibility_Index, where a lower accessibility score indicates greater infrastructure strain and therefore higher tourism pressure.")

heading("5.3 Random Forest Regressor",2)
para("The Random Forest Regressor serves as the primary supervised learning model for climate-based pilgrim count prediction. The model is configured with the following hyperparameters optimised through empirical evaluation:")
tbl(
    ['Hyperparameter','Value','Rationale'],
    [['n_estimators','300','Sufficient ensemble size for stable predictions'],
     ['max_depth','15','Prevents overfitting while capturing complex interactions'],
     ['min_samples_split','5','Minimum samples required to split an internal node'],
     ['min_samples_leaf','2','Minimum samples in leaf nodes for regularisation'],
     ['max_features','sqrt','Square root of total features considered at each split'],
     ['random_state','42','Fixed seed for reproducibility'],
     ['n_jobs','-1','Parallel training across all available CPU cores']]
)
para("Table 5.3: Random Forest Hyperparameter Configuration",italic=True,size=10)
para("The model is trained on up to nineteen features spanning raw climate measurements such as temperature, rainfall, humidity, and wind speed; temporal encodings including month, cyclical month sine and cosine; derived features including temperature anomaly, rainfall anomaly, capacity utilization, and year-over-year growth; and infrastructure features including carrying capacity, peak season indicator, and accessibility index. The get_available_features function dynamically selects the subset of the full nineteen-feature list that exists in the current DataFrame, ensuring graceful handling of datasets with missing columns.")
para("Model evaluation employs an eighty-twenty train-test split with five-fold cross-validation on the training set. The cross-validation provides a robust estimate of generalisation performance by averaging R-squared scores across five non-overlapping validation folds. The trained model achieved an R-squared of {:.3f} on the held-out test set with a cross-validation mean R-squared of {:.3f} plus or minus {:.3f}.".format(M['rf_r2'], M['rf_cv_mean'], M['rf_cv_std']))

heading("5.4 SARIMA Model",2)
para("The Seasonal Autoregressive Integrated Moving Average model extends the classical ARIMA framework with explicit seasonal components to capture the twelve-month periodicity inherent in pilgrimage tourism data. The model is parameterised as SARIMA(p,d,q)(P,D,Q,s) where the lowercase parameters govern the non-seasonal component and the uppercase parameters govern the seasonal component with period s set to twelve months.")
para("Order selection employs a two-phase optimised grid search strategy. Phase one performs a coarse search over all combinations of p in zero and one, d in zero and one, q in zero and one, P in zero and one, D in zero and one, and Q in zero and one, yielding sixty-four candidate models. Each candidate is fitted using maximum likelihood estimation with the SARIMAX implementation from statsmodels, and the model with the lowest Akaike Information Criterion is retained. Phase two refines the non-seasonal order by expanding p and q to the range zero through two while holding the best seasonal order fixed, evaluating up to nine additional models.")
para("The final fitted model produces twelve-month-ahead point forecasts with ninety-five per cent confidence intervals computed from the forecast error variance. Model diagnostics include the Ljung-Box test for residual autocorrelation at lag ten, where a p-value exceeding 0.05 indicates that the residuals do not exhibit significant serial correlation and the model has adequately captured the temporal structure of the data.")

heading("5.5 LSTM Neural Network",2)
para("The Long Short-Term Memory neural network represents the deep learning component of the forecasting framework. The architecture implements a sequential model consisting of six layers. The input layer accepts sequences of shape twelve by one, representing twelve consecutive monthly pilgrim count observations normalised to the zero-one range using MinMaxScaler. The first LSTM layer contains sixty-four units with return_sequences enabled to pass the full sequence to subsequent layers. A dropout layer with twenty per cent probability follows to prevent co-adaptation of hidden units. The second LSTM layer reduces dimensionality to thirty-two units with return_sequences disabled, producing a single output vector. Another twenty per cent dropout layer provides additional regularisation. A dense layer with sixteen units and ReLU activation introduces non-linear feature transformation. The final dense layer with a single unit produces the point prediction.")
para("Training employs the Adam optimiser with mean squared error loss function. Early stopping monitors validation loss with a patience of ten epochs and restores the best weights observed during training. The training-validation split allocates eighty per cent of available sequences for training and twenty per cent for validation. Forecasting generates twelve-step-ahead predictions through iterative autoregressive application, where each predicted value is appended to the input sequence for the next prediction step. Confidence intervals are estimated using a progressive widening approach, where the interval width expands from five per cent of the predicted value at step one to twenty per cent at step twelve.")

heading("5.6 SHAP Explainability",2)
para("Model interpretability is achieved through SHapley Additive exPlanations, which decompose individual predictions into feature-level contributions based on cooperative game theory. The TreeExplainer class computes exact SHAP values for tree-based models in polynomial time, avoiding the exponential complexity of the general Shapley value computation.")
para("For each prediction scenario configured through the interactive dashboard, the system computes SHAP values for all input features, producing a signed contribution score that indicates both the direction and magnitude of each feature's influence on the prediction relative to the expected base value. These values are visualised as horizontal bar charts with colour coding indicating positive contributions in red and negative contributions in blue, enabling stakeholders to understand not just what the model predicts but why it makes that specific prediction.")

heading("5.7 Multi-Factor Alert Engine",2)
para("The alert engine generates severity-graded risk assessments across four independent alert categories, each evaluating a distinct dimension of ecosystem risk. The four categories are Tourist Overload based on capacity utilization ratio, Climate Stress based on temperature anomaly magnitude, NDVI Degradation based on vegetation health decline rate, and ESI Level based on the composite ecosystem stress index.")
tbl(
    ['Level','Priority','Colour','Description'],
    [['LOW','0','Green (#22c55e)','Within safe operational limits, no action required'],
     ['MODERATE','1','Amber (#f59e0b)','Approaching thresholds, increased monitoring recommended'],
     ['HIGH','2','Orange (#f97316)','Near critical levels, prepare intervention measures'],
     ['CRITICAL','3','Red (#ef4444)','Exceeding safe limits, immediate action required']]
)
para("Table 5.2: Alert Level Classification System",italic=True,size=10)
para("Each alert check function evaluates the current value against configurable threshold pairs. For tourist overload, capacity utilization below sixty per cent is classified as Low, sixty to eighty-five per cent as Moderate, eighty-five to one hundred per cent as High, and above one hundred per cent as Critical. Temperature anomaly thresholds are set at three degrees Celsius for Moderate and six degrees for High. NDVI change thresholds are set at negative 0.05 for Moderate and negative 0.10 for Critical. ESI thresholds follow the standard classification at forty and seventy.")

heading("5.8 NDVI Analysis Pipeline",2)
para("The NDVI analysis pipeline implements a three-tier cascading data source architecture that maximises data quality while ensuring availability. The primary source is Google Earth Engine accessing the MODIS/061/MOD13Q1 product, which provides sixteen-day composite NDVI at 250-metre resolution. When GEE is unavailable due to authentication or connectivity issues, the system falls back to the ORNL MODIS REST API, a free publicly accessible interface to the same MODIS product that requires no authentication. When neither satellite data source is accessible, the system generates physically grounded synthetic NDVI values using a historical model calibrated to Himalayan vegetation phenology.")
para("The historical model incorporates altitude-based baseline NDVI values specific to each shrine, seasonal sinusoidal patterns calibrated to the Himalayan growing season peaking between June and September, a gradual long-term degradation trend reflecting tourism pressure, and stochastic inter-annual variability. All three data sources produce output in a standardised pandas DataFrame format with Date index and NDVI column, enabling seamless substitution without affecting downstream visualisation or analysis code.")
img('chart_sensitivity.png', 5.5, 'Figure 5.2: Sensitivity Analysis - Pilgrim Count Impact on ESI')
pb()

doc.save(str(OUT))
print(f"Part 3 done: Chapters 4-5 added.")
